import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from groq import Groq
import io
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Page config with modern dark theme
st.set_page_config(
    page_title="AI Business Dashboard",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Theme toggle in session state
if "theme" not in st.session_state:
    st.session_state.theme = "dark"

# Initialize Groq API client
try:
    client = Groq(api_key=st.secrets["GROQ_API_KEY"])
    model_ready = True
except Exception as e:
    st.error(f"❌ API Configuration Error: {e}")
    model_ready = False

# Define color schemes
THEMES = {
    "dark": {
        "bg_gradient": "linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #1f2b3e 100%)",
        "sidebar_bg": "linear-gradient(180deg, #1f2b3e 0%, #1a1a2e 100%)",
        "accent": "#7c6af7",
        "accent_secondary": "#f472b6",
        "accent_tertiary": "#34d399",
        "text_primary": "#f0f0f5",
        "text_secondary": "#c8cfe0",
        "text_tertiary": "#a8b4c8",
        "success": "#34d399",
        "warning": "#fbbf24",
        "error": "#f87171",
        "info": "#7c6af7",
        "metric_bg": "#2a2550",
        "button_bg": "#7c6af7",
        "button_hover": "#6350e8",
        "container_bg": "#1f2b3e",
        "scrollbar_primary": "#7c6af7",
        "scrollbar_hover": "#f472b6"
    },
    "light": {
        "bg_gradient": "linear-gradient(135deg, #fdf6ff 0%, #f3f0ff 100%)",
        "sidebar_bg": "linear-gradient(180deg, #ede9fe 0%, #f3f0ff 100%)",
        "accent": "#6d56f5",
        "accent_secondary": "#ec4899",
        "accent_tertiary": "#10b981",
        "text_primary": "#1e1b4b",
        "text_secondary": "#3b3573",
        "text_tertiary": "#5b5499",
        "success": "#10b981",
        "warning": "#f59e0b",
        "error": "#ef4444",
        "info": "#6d56f5",
        "metric_bg": "#ede9fe",
        "button_bg": "#6d56f5",
        "button_hover": "#5840e0",
        "container_bg": "#f3f0ff",
        "scrollbar_primary": "#6d56f5",
        "scrollbar_hover": "#ec4899"
    }
}

theme = THEMES[st.session_state.theme]

# AI content text: pure white in dark mode, theme primary in light mode
ai_text_color = "#ffffff" if st.session_state.theme == "dark" else theme["text_primary"]

# Custom CSS with dynamic theme
css_content = f"""
<style>
    * {{
        color-scheme: {'dark' if st.session_state.theme == 'dark' else 'light'};
    }}
    
    body {{
        background: {theme['bg_gradient']};
        color: {theme['text_primary']};
    }}
    
    .stApp {{
        background: {theme['bg_gradient']};
    }}
    
    /* Main container */
    .stMainBlockContainer {{
        padding: 40px 20px 20px 20px;
    }}
    
    /* Headers with vibrant gradients */
    h1 {{
        background: linear-gradient(90deg, {theme['accent']}, {theme['accent_secondary']}, {theme['accent_tertiary']});
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        font-weight: 900;
        font-size: 2.5em !important;
        padding-top: 24px !important;
        margin-top: 16px !important;
    }}
    
    h2, h3 {{
        background: linear-gradient(90deg, {theme['success']}, {theme['accent']}, {theme['accent_secondary']});
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        font-weight: 800;
    }}
    
    /* Restore emoji natural colours inside gradient headings */
    h1 span[role="img"], h2 span[role="img"], h3 span[role="img"],
    h1 .stEmoji, h2 .stEmoji, h3 .stEmoji {{
        -webkit-text-fill-color: initial !important;
        background: none !important;
        background-clip: initial !important;
        filter: none !important;
    }}
    
    /* Sidebar styling */
    .stSidebar {{
        background: {theme['sidebar_bg']};
        border-right: 2px solid {theme['accent']};
    }}
    
    .stSidebar [data-testid="stSidebarNav"] {{
        background: transparent;
    }}
    
    /* Radio buttons and selection */
    [data-testid="stRadio"] {{
        background-color: {theme['container_bg']};
        padding: 15px;
        border-radius: 12px;
        border-left: 4px solid {theme['accent']};
    }}
    
    /* Metrics */
    .stMetric {{
        background: {theme['metric_bg']};
        padding: 20px;
        border-radius: 12px;
        border-left: 4px solid {theme['accent']};
        box-shadow: 0px 8px 20px rgba(124, 106, 247, 0.2);
    }}
    
    .stMetric label {{
        color: {theme['text_primary']} !important;
        font-weight: 700 !important;
    }}
    
    .stMetric [data-testid="stMetricValue"] {{
        color: {theme['accent']} !important;
        font-size: 24px !important;
        font-weight: 700 !important;
    }}
    
    /* Buttons */
    .stButton > button {{
        background: linear-gradient(90deg, {theme['button_bg']}, {theme['button_hover']});
        color: {theme['text_primary']} !important;
        border: none;
        border-radius: 10px;
        padding: 12px 28px;
        font-weight: 700;
        font-size: 16px;
        box-shadow: 0px 8px 20px rgba(124, 106, 247, 0.35);
        transition: all 0.3s ease;
        width: 100%;
    }}
    
    .stButton > button:hover {{
        background: linear-gradient(90deg, {theme['button_hover']}, {theme['accent_secondary']});
        box-shadow: 0px 12px 30px rgba(124, 106, 247, 0.5);
        transform: translateY(-2px);
    }}
    
    /* Selectbox and input fields */
    .stSelectbox, .stMultiSelect, .stFileUploader {{
        background-color: {theme['container_bg']} !important;
    }}
    
    .stSelectbox [data-baseweb="select"] {{
        background: {theme['container_bg']} !important;
        border: 2px solid {theme['accent']} !important;
        border-radius: 8px !important;
        color: {theme['text_primary']} !important;
    }}
    
    .stMultiSelect [data-baseweb="select"] {{
        background: {theme['container_bg']} !important;
        border: 2px solid {theme['success']} !important;
        border-radius: 8px !important;
        color: {theme['text_primary']} !important;
    }}
    
    /* Tabs */
    .stTabs [data-baseweb="tab-list"] {{
        background-color: {theme['container_bg']};
        border-bottom: 2px solid {theme['accent']};
    }}
    
    .stTabs [data-baseweb="tab"] {{
        background-color: {theme['container_bg']};
        border-radius: 8px 8px 0 0;
        color: {theme['text_tertiary']};
    }}
    
    .stTabs [aria-selected="true"] {{
        background: linear-gradient(90deg, {theme['accent']}, {theme['button_hover']});
        color: {theme['text_primary']};
    }}
    
    /* Alert messages */
    .stSuccess {{
        background: {'linear-gradient(135deg, #0d3b2e 0%, #065f46 100%)' if st.session_state.theme == 'dark' else 'linear-gradient(135deg, #d1fae5 0%, #a7f3d0 100%)'};
        border-left: 4px solid {theme['success']};
        border-radius: 8px;
        padding: 15px;
        color: {theme['text_primary']} !important;
    }}
    
    .stWarning {{
        background: {'linear-gradient(135deg, #3d2a00 0%, #78460a 100%)' if st.session_state.theme == 'dark' else 'linear-gradient(135deg, #fef3c7 0%, #fde68a 100%)'};
        border-left: 4px solid {theme['warning']};
        border-radius: 8px;
        padding: 15px;
        color: {theme['text_primary']} !important;
    }}
    
    .stError {{
        background: {'linear-gradient(135deg, #3b0d0d 0%, #7f1d1d 100%)' if st.session_state.theme == 'dark' else 'linear-gradient(135deg, #fee2e2 0%, #fecaca 100%)'};
        border-left: 4px solid {theme['error']};
        border-radius: 8px;
        padding: 15px;
        color: {theme['text_primary']} !important;
    }}
    
    .stInfo {{
        background: {'linear-gradient(135deg, #1e1250 0%, #2d1f8a 100%)' if st.session_state.theme == 'dark' else 'linear-gradient(135deg, #ede9fe 0%, #ddd6fe 100%)'};
        border-left: 4px solid {theme['info']};
        border-radius: 8px;
        padding: 15px;
        color: {theme['text_primary']} !important;
    }}
    
    /* Divider */
    .stDivider {{
        border-color: {theme['accent']};
        opacity: 0.6;
    }}
    
    /* Captions */
    .stCaption {{
        color: {theme['accent']} !important;
        font-weight: 500;
    }}
    
    /* Global paragraph colour */
    p {{
        color: {theme['text_primary']} !important;
    }}
    
    /* File uploader */
    [data-testid="stFileUploader"] button {{
        color: #ffffff !important;
        background: linear-gradient(90deg, {theme['button_bg']}, {theme['button_hover']}) !important;
        border: none !important;
        font-weight: 600 !important;
    }}
    [data-testid="stFileUploader"] button span {{
        color: #ffffff !important;
    }}
    [data-testid="stFileUploader"] label {{
        color: {theme['text_primary']} !important;
    }}
    [data-testid="stFileUploaderDropzone"] {{
        border: 2px dashed {theme['accent']} !important;
        border-radius: 10px !important;
        background: {theme['container_bg']} !important;
    }}
    [data-testid="stFileUploaderDropzone"] span,
    [data-testid="stFileUploaderDropzone"] p {{
        color: {theme['text_primary']} !important;
    }}
    
    /* Scrollbar */
    ::-webkit-scrollbar {{ width: 12px; height: 12px; }}
    ::-webkit-scrollbar-track {{ background: {theme['container_bg']}; }}
    ::-webkit-scrollbar-thumb {{ background: {theme['scrollbar_primary']}; border-radius: 6px; }}
    ::-webkit-scrollbar-thumb:hover {{ background: {theme['scrollbar_hover']}; }}
    
    /* Expander */
    .stExpander {{
        background-color: {theme['container_bg']};
        border: 1px solid {theme['accent']};
        border-radius: 8px;
    }}
    
    /* DataFrames */
    .stDataFrame {{ background-color: {theme['container_bg']}; }}
    .stDataFrame thead {{ background-color: {theme['accent']} !important; color: {theme['text_primary']} !important; }}
    .stDataFrame tbody tr {{ border-bottom: 1px solid {theme['accent']}; }}
    .stDataFrame tbody tr:hover {{ background-color: {theme['metric_bg']}; }}
    .stDataFrame {{ color: {theme['text_primary']} !important; }}
    
    /* ── Analysis container: force ALL child text to white (dark) / primary (light) ── */
    .analysis-container {{
        background: {theme['container_bg']};
        padding: 24px;
        border-radius: 12px;
        border-left: 4px solid {theme['accent']};
        box-shadow: 0px 8px 20px rgba(124, 106, 247, 0.2);
    }}
    
    .analysis-container *,
    .analysis-container p,
    .analysis-container span,
    .analysis-container div,
    .analysis-container ul,
    .analysis-container ol,
    .analysis-container li,
    .analysis-container strong,
    .analysis-container b,
    .analysis-container em,
    .analysis-container i,
    .analysis-container a {{
        color: {ai_text_color} !important;
    }}
    
    /* Section headings inside AI output — accent colour, no gradient */
    .analysis-container h1,
    .analysis-container h2,
    .analysis-container h3,
    .analysis-container h4 {{
        color: {theme['accent_tertiary']} !important;
        -webkit-text-fill-color: {theme['accent_tertiary']} !important;
        background: none !important;
        margin-top: 18px;
        margin-bottom: 8px;
        font-weight: 700;
    }}
</style>
"""

st.markdown(css_content, unsafe_allow_html=True)


# ── Helper: build .docx bytes from AI markdown text ───────────────────────────
def _add_runs_with_bold(paragraph, text: str):
    """Split on **bold** markers and add styled runs."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def generate_docx(markdown_text: str, title: str = "AI Business Insights Report") -> bytes:
    doc = DocxDocument()

    # Document title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title_para.runs:
        run = title_para.runs[0]
        run.font.color.rgb = RGBColor(0x6D, 0x56, 0xF5)
        run.font.size = Pt(22)

    doc.add_paragraph()  # spacer

    for line in markdown_text.splitlines():
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith("### "):
            h = doc.add_heading(stripped[4:].strip(), level=3)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)

        elif stripped.startswith("## "):
            h = doc.add_heading(stripped[3:].strip(), level=2)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x7C, 0x6A, 0xF7)

        elif stripped.startswith("# "):
            h = doc.add_heading(stripped[2:].strip(), level=1)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x7C, 0x6A, 0xF7)

        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(p, stripped[2:].strip())

        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_bold(p, re.sub(r"^\d+\.\s", "", stripped))

        elif stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
            run.font.color.rgb = RGBColor(0x6D, 0x56, 0xF5)
            run.font.size = Pt(12)

        else:
            p = doc.add_paragraph()
            _add_runs_with_bold(p, stripped)
            p.paragraph_format.space_after = Pt(4)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── Sidebar ────────────────────────────────────────────────────────────────────
col1, col2 = st.sidebar.columns([0.8, 0.2])
with col1:
    st.sidebar.title("📊 Dashboard")
with col2:
    theme_icon = "🌙" if st.session_state.theme == "dark" else "☀️"
    if st.sidebar.button(theme_icon, help="Toggle Light/Dark Mode", key="theme_toggle"):
        st.session_state.theme = "light" if st.session_state.theme == "dark" else "dark"
        st.rerun()

st.sidebar.markdown("---")
menu = st.sidebar.radio(
    "Navigation",
    ["📤 Upload Data", "📈 Dashboard", "🤖 AI Insights"],
    help="Select a section to navigate"
)

# ── Main title ─────────────────────────────────────────────────────────────────
st.title("📊 AI Business Dashboard")
st.caption("✨ Transform your data into powerful business insights using AI")

# ── Upload section ─────────────────────────────────────────────────────────────
if menu == "📤 Upload Data":
    st.subheader("📤 Upload Your Data")
    col1, col2 = st.columns([3, 1])
    
    with col1:
        file = st.file_uploader(
            "📁 Upload your CSV file",
            type=["csv"],
            help="Select a CSV file to analyze"
        )
    
    if file:
        try:
            df = pd.read_csv(file)
            df.to_csv("data.csv", index=False)
            st.success("✅ File uploaded successfully!")
            st.info(f"📊 Dataset: {len(df)} rows × {len(df.columns)} columns")
            st.subheader("📋 Data Preview")
            st.dataframe(df.head(10), use_container_width=True)
        except Exception as e:
            st.error(f"❌ Error uploading file: {e}")
else:
    file = None

# ── Load data ──────────────────────────────────────────────────────────────────
try:
    df = pd.read_csv("data.csv")
except FileNotFoundError:
    df = None
except Exception as e:
    st.error(f"❌ Error loading data: {e}")
    df = None

# ── Dashboard ──────────────────────────────────────────────────────────────────
if menu == "📈 Dashboard" and df is not None:
    st.sidebar.header("🔍 Filters")
    df_filtered = df.copy()
    cat_cols = df.select_dtypes(include='object').columns

    for col in cat_cols:
        options = df[col].unique()
        selected = st.sidebar.multiselect(f"Filter {col}", options)
        if selected:
            df_filtered = df_filtered[df_filtered[col].isin(selected)]

    st.subheader("📊 Key Metrics")
    num_cols = df_filtered.select_dtypes(include='number').columns

    if len(num_cols) > 0:
        selected_kpi = st.selectbox("📈 Select KPI Column", num_cols)
        total    = df_filtered[selected_kpi].sum()
        avg      = df_filtered[selected_kpi].mean()
        max_val  = df_filtered[selected_kpi].max()
        min_val  = df_filtered[selected_kpi].min()

        c1, c2, c3, c4 = st.columns(4, gap="large")
        with c1: st.metric("💰 Total",    f"{total:,.2f}")
        with c2: st.metric("📈 Average",  f"{avg:,.2f}")
        with c3: st.metric("⬆️ Maximum",  f"{max_val:,.2f}")
        with c4: st.metric("⬇️ Minimum",  f"{min_val:,.2f}")

    st.divider()
    st.subheader("📈 Visual Analysis")

    chart_type    = st.selectbox("📊 Choose Chart Type",
                                 ["Bar Chart", "Line Chart", "Pie Chart", "Scatter Plot", "Area Chart"])
    plot_template = "plotly_dark" if st.session_state.theme == "dark" else "plotly"
    bg_color      = theme['container_bg'] if st.session_state.theme == "dark" else "#ffffff"

    def chart_layout():
        return dict(plot_bgcolor=bg_color, paper_bgcolor=bg_color,
                    font=dict(color=theme['text_primary'], size=12))

    if chart_type == "Bar Chart" and len(cat_cols) > 0 and len(num_cols) > 0:
        x_col = st.selectbox("X-axis", cat_cols, key="bar_x")
        y_col = st.selectbox("Y-axis", num_cols, key="bar_y")
        fig = px.bar(df_filtered, x=x_col, y=y_col,
                     color_discrete_sequence=[theme['accent']], template=plot_template)
        fig.update_layout(**chart_layout(), hovermode='x unified')
        st.plotly_chart(fig, use_container_width=True)

    elif chart_type == "Line Chart" and len(num_cols) > 0:
        y_col = st.selectbox("Y-axis", num_cols, key="line_y")
        fig = px.line(df_filtered, y=y_col,
                      color_discrete_sequence=[theme['success']], template=plot_template)
        fig.update_layout(**chart_layout(), hovermode='x unified')
        st.plotly_chart(fig, use_container_width=True)

    elif chart_type == "Pie Chart" and len(cat_cols) > 0 and len(num_cols) > 0:
        names  = st.selectbox("Category", cat_cols, key="pie_names")
        values = st.selectbox("Values",   num_cols, key="pie_values")
        fig = px.pie(df_filtered, names=names, values=values,
                     color_discrete_sequence=[theme['accent'], theme['success'], theme['accent_tertiary'],
                                              theme['warning'], theme['error'], theme['accent_secondary'], theme['info']],
                     template=plot_template)
        fig.update_layout(**chart_layout())
        st.plotly_chart(fig, use_container_width=True)

    elif chart_type == "Scatter Plot" and len(num_cols) > 1:
        x_col = st.selectbox("X-axis", num_cols, key="scatter_x")
        y_col = st.selectbox("Y-axis", num_cols, key="scatter_y")
        fig = px.scatter(df_filtered, x=x_col, y=y_col,
                         color_discrete_sequence=[theme['accent_tertiary']], template=plot_template)
        fig.update_layout(**chart_layout(), hovermode='closest')
        st.plotly_chart(fig, use_container_width=True)

    elif chart_type == "Area Chart" and len(num_cols) > 0:
        y_col = st.selectbox("Y-axis", num_cols, key="area_y")
        fig = px.area(df_filtered, y=y_col,
                      color_discrete_sequence=[theme['warning']], template=plot_template)
        fig.update_layout(**chart_layout(), hovermode='x unified')
        st.plotly_chart(fig, use_container_width=True)

    st.divider()
    st.subheader("🏆 Top Performers")

    if len(cat_cols) > 0 and len(num_cols) > 0:
        cat      = cat_cols[0]
        num      = num_cols[0]
        top_data = df_filtered.groupby(cat)[num].sum().sort_values(ascending=False).head(5)
        fig = px.bar(x=top_data.index, y=top_data.values,
                     labels={"x": cat, "y": num},
                     color_discrete_sequence=[theme['success']], template=plot_template)
        fig.update_layout(**chart_layout(), xaxis_title=cat, yaxis_title=num, hovermode='x unified')
        st.plotly_chart(fig, use_container_width=True)

# ── AI Insights ────────────────────────────────────────────────────────────────
if menu == "🤖 AI Insights" and df is not None:
    st.subheader("🤖 AI Insights Engine - Powered by Groq ⚡")
    st.info("💡 Lightning-fast AI analysis using Groq's latest Llama models")

    if not model_ready:
        st.error("❌ Groq API is not configured. Please add GROQ_API_KEY to your Streamlit secrets.")
        st.write("**How to fix (3 easy steps):**")
        c1, c2 = st.columns(2)
        with c1:
            st.subheader("🌐 Streamlit Cloud")
            st.write("1. Go to **Manage app** (bottom right)")
            st.write("2. Click **Secrets**")
            st.write("3. Add this:")
            st.code("GROQ_API_KEY = 'your-key-here'")
        with c2:
            st.subheader("💻 Local Development")
            st.write("1. Create `.streamlit/secrets.toml`")
            st.write("2. Add this:")
            st.code("GROQ_API_KEY = 'your-key-here'")
            st.write("3. Run: `streamlit run app.py`")
        st.write("**Get Free API Key (25 req/min):**")
        st.write("👉 https://console.groq.com")

    else:
        if st.button("🚀 Generate AI Insights", use_container_width=True):
            with st.spinner("⚡ Groq is analyzing your data (lightning fast)..."):
                try:
                    data_summary = df.describe().to_string()

                    prompt = f"""You are a professional business analyst with expertise in data analysis and strategic insights.

Analyze this dataset and provide comprehensive business intelligence:

DATASET STATISTICS:
{data_summary}

DATASET INFORMATION:
- Total Records: {len(df):,}
- Total Columns: {len(df.columns)}
- Data Columns: {', '.join(df.columns)}
- Data Types: {dict(df.dtypes)}

Please provide a detailed analysis with these sections:

1. **📊 Executive Summary**
   - What does this data represent?
   - Overall health and status of the data

2. **📈 Key Trends & Patterns**
   - Identify main trends and patterns
   - Growth or decline indicators
   - Seasonal patterns (if applicable)

3. **💡 Business Insights**
   - Most important findings
   - What the metrics tell about business performance
   - Correlations and relationships

4. **⚠️ Risks & Concerns**
   - Potential risks or red flags
   - Areas of concern
   - Issues that need attention

5. **🎯 Opportunities**
   - Growth opportunities
   - Areas for improvement
   - Untapped potential

6. **💼 Recommendations**
   - Top 3-5 actionable recommendations
   - Priority order
   - Expected impact

Please format with clear headings and bullet points for easy reading."""

                    response = client.chat.completions.create(
                        model="llama-3.1-8b-instant",
                        messages=[{"role": "user", "content": prompt}],
                        max_tokens=2048
                    )

                    ai_content = response.choices[0].message.content

                    # Persist for download button
                    st.session_state["last_ai_report"] = ai_content

                    st.success("✅ Analysis Complete!")

                    # ── Render AI content with forced white text in dark mode ──
                    st.markdown(f"""
                    <style>
                    .analysis-container {{
                        background: {theme['container_bg']};
                        padding: 24px;
                        border-radius: 12px;
                        border-left: 4px solid {theme['accent']};
                        box-shadow: 0px 8px 20px rgba(124, 106, 247, 0.2);
                    }}
                    .analysis-container *,
                    .analysis-container p,
                    .analysis-container span,
                    .analysis-container div,
                    .analysis-container ul,
                    .analysis-container ol,
                    .analysis-container li,
                    .analysis-container strong,
                    .analysis-container b,
                    .analysis-container em,
                    .analysis-container i,
                    .analysis-container a {{
                        color: {ai_text_color} !important;
                    }}
                    .analysis-container h1,
                    .analysis-container h2,
                    .analysis-container h3,
                    .analysis-container h4 {{
                        color: {theme['accent_tertiary']} !important;
                        -webkit-text-fill-color: {theme['accent_tertiary']} !important;
                        background: none !important;
                        margin-top: 18px;
                        margin-bottom: 8px;
                        font-weight: 700;
                    }}
                    </style>
                    <div class="analysis-container">{ai_content}</div>
                    """, unsafe_allow_html=True)

                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")
                    st.info("**Troubleshooting:**")
                    st.write("- ✓ API key is correct?")
                    st.write("- ✓ Rate limit (25 req/min) exceeded?")
                    st.write("- ✓ Internet connection stable?")
                    st.write("- ✓ Try with smaller dataset?")

        # ── Download report button (appears once report is generated) ──────────
        if st.session_state.get("last_ai_report"):
            st.divider()
            st.markdown("### 📥 Download Report")
            docx_bytes = generate_docx(st.session_state["last_ai_report"])
            st.download_button(
                label="📄 Download as Word Document (.docx)",
                data=docx_bytes,
                file_name="AI_Business_Insights_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

# ── No data state ──────────────────────────────────────────────────────────────
if df is None:
    st.warning("⚠️ No data loaded. Please upload a dataset to get started.")
    st.info("📝 Navigate to '📤 Upload Data' section to upload your CSV file.")
